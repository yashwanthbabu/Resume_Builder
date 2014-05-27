from django.http.response import Http404
from django.shortcuts import render, redirect
from django.http import HttpResponseRedirect, HttpResponse
from res1.models import *
from django.shortcuts import render_to_response
#from flask import Flask, render_template, redirect, url_for
from django.core.mail import send_mail,EmailMessage,BadHeaderError
from pdfs import pdf
import cStringIO as StringIO
import ho.pisa as pisa
#from generic_mail import Email
from django.template.loader import get_template,render_to_string

from django.template import loader, Context
from django.utils.translation import ugettext as _
from django.core.exceptions import ValidationError
from docx import Document


def home(request):
    return render(request,'home.html',locals())

def registration(request):
    if request.is_ajax() and request.method == "POST":
        global email
	fullname = request.POST.get("fullname")
	email = request.POST.get("emailid")
	if email and Name.objects.filter(Email_id=email):
            return render(request,'email.html')
        #  Implement save function in mongodb
	#import pdb;pdb.set_trace()
	nam = Name.objects.create(Full_name=fullname,Email_id=email)
        print nam
        id = nam.id
        print id
        return render(request,'education.html',{'objects_id':nam.id,'mail':email,'name':fullname})
	#nme = Name.Objects.all()
    else:
	raise Http404
    return render(request,'home.html',{'mail':email,'name':fullname})

   #return render(request,'education.html',{'objects':nam})
def addmore(request):
    if request.is_ajax() and request.method == "POST":
	quafn = request.POST.get("quafn")
	univ = request.POST.get("univ")
	year = request.POST.get("year")
	perc = request.POST.get("perc")
	#  Implement save function in mongodb
        #import pdb;pdb.set_trace()
	edu = Education.objects.create(Qualification=quafn,University=univ,Year_of_Passing=year,Aggrigation=perc)
	# after save implemented
        obj_id = request.POST.get("y1")
        print obj_id
        name_obj = Name.objects.get(id=obj_id)
        print name_obj
        edu.edn = name_obj
        edu.save()
	m_id = request.POST.get("m1")
        print m_id
	n_id = request.POST.get("n1")
        print n_id
        print request
	cmp = request.POST.get("cmp",None)
	if cmp == "true":
	    return render(request,'professional.html',{'objects_id':obj_id,'mail':m_id,'name':n_id,'qua':quafn,'univ':univ,'year':year,'perc':perc})
    else:
	raise Http404
    return render(request,'education.html',{'objects_id':obj_id,'mail':m_id,'name':n_id,'qua':quafn,'univ':univ,'year':year,'perc':perc})
    


def profn(request):
    
    if request.is_ajax() and request.method == "POST":
	job = request.POST.get("job",None)
	exp = request.POST.get("exp",None)
	ind = request.POST.get("ind",None)
	func = request.POST.get("func",None)
	key = request.POST.get("key",None)
	#  Implement save function in mongodb
	pro = Profession.objects.create(Job_Type=job, Total_Experience_in_months=exp,Current_Industry=ind,Functional_Area=func,Key_Skills=key)
        obj_id = request.POST.get("y2")
        print obj_id
        m_id = request.POST.get("m2")
        print m_id
        n_id = request.POST.get("n2")
        print n_id
        q_id = request.POST.get("q1")
        print q_id
        u_id = request.POST.get("u1")
        print u_id
        y_id = request.POST.get("y11")
        print y_id
        p_id = request.POST.get("p1")
        print p_id
        print request
        try:
            nam = Name.objects.get(id=obj_id)
            print nam
        except Name.DoesNotExist:
            nam = None
        pro.prof = nam
        pro.save()
	cmp = request.POST.get("cmp",None)
	if cmp == "true":
	    return render(request,'contact.html',{'objects_id':obj_id,'mail':m_id,'name':n_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':job,'exp':exp,'ind':ind,'func':func,'key':key})
        else:
	    raise Http404
    return render(request,'professional.html',{'objects_id':obj_id,'mail':m_id,'name':n_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':job,'exp':exp,'ind':ind,'func':func,'key':key})


def send(request):
    
    username = "yyyy"
    print username
    message = "hi"
    subject = "Confirmation mail"
    from_email = 'test@creyo.com'
    a = Name.objects.all()
    print a
    b = Education.objects.all()
    print b
    c = Profession.objects.all()
    print c
    z_id = request.POST.get("z1")
    print z_id
    m_id = request.POST.get("m3")
    print m_id
    n_id = request.POST.get("n3")
    print n_id
    q_id = request.POST.get("q2")
    print q_id
    u_id = request.POST.get("u2")
    print u_id
    y_id = request.POST.get("y12")
    print y_id
    p_id = request.POST.get("p2")
    print p_id
    j_id = request.POST.get("j1")
    print j_id
    e_id = request.POST.get("e1")
    print e_id
    i_id = request.POST.get("i1")
    print i_id
    f_id = request.POST.get("f1")
    print f_id
    k_id = request.POST.get("k1")
    print k_id
    print request
    to = [m_id]
    #message = get_template('/home/naveen/Resume1/res1/templates/yashwanth1')
    msg = get_template('/home/naveen/Resume1/res1/templates/c.txt').render(Context({'fullname':n_id,'mailing':m_id}))
    radio = request.POST.get('imgControl')
    print type(radio),'yashu',radio
    if int(radio) == 0:
        print 'radio'
        message = get_template('/home/naveen/Resume1/res1/templates/Format1.html').render(Context({'fullname':n_id,'mailing':m_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':j_id,'exp':e_id,'ind':i_id,'func':f_id,'key':k_id}))
        #result = StringIO.StringIO()
        result = StringIO.StringIO()
        result.writelines(message)
        #pdf = pisa.pisaDocument(StringIO.StringIO(message), result)
        try:
        
            mail = EmailMessage(subject,msg,from_email,to)
            #attachment = myview(request)
            mail.attach('Resume.docx',result.getvalue(),'application/ms-word')
            mail.send(msg)
            return render(request,'text.html',{'mail':m_id,'name':n_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':j_id,'exp':e_id,'ind':i_id,'func':f_id,'key':k_id})
        except BadHeaderError:
            return HttpResponse('Invalid header found.')
        return render(request,'text.html',{'objects_id':z_id,'mail':m_id,'name':n_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':j_id,'exp':e_id,'ind':i_id,'func':f_id,'key':k_id})
    
    elif int(radio) == 1:
        print 'radio'
        message = get_template('/home/naveen/Resume1/res1/templates/Format2.html').render(Context({'fullname':n_id,'mailing':m_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':j_id,'exp':e_id,'ind':i_id,'func':f_id,'key':k_id}))
        #result = StringIO.StringIO()
        result = StringIO.StringIO()
        result.writelines(message)
        #pdf = pisa.pisaDocument(StringIO.StringIO(message), result)
        try:
        
            mail = EmailMessage(subject,msg,from_email,to)
            #attachment = myview(request)
            mail.attach('Resume.docx',result.getvalue(),'application/ms-word')
            mail.send(msg)
            return render(request,'text.html',{'mail':m_id,'name':n_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':j_id,'exp':e_id,'ind':i_id,'func':f_id,'key':k_id})
        except BadHeaderError:
            return HttpResponse('Invalid header found.')
        return render(request,'text.html',{'objects_id':z_id,'mail':m_id,'name':n_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':j_id,'exp':e_id,'ind':i_id,'func':f_id,'key':k_id})

    elif int(radio) == 2:
        print 'radio'
        message = get_template('/home/naveen/Resume1/res1/templates/Format3.html').render(Context({'fullname':n_id,'mailing':m_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':j_id,'exp':e_id,'ind':i_id,'func':f_id,'key':k_id}))
        #result = StringIO.StringIO()
        result = StringIO.StringIO()
        result.writelines(message)
        #pdf = pisa.pisaDocument(StringIO.StringIO(message), result)
        try:
        
            mail = EmailMessage(subject,msg,from_email,to)
            #attachment = myview(request)
            mail.attach('Resume.docx',result.getvalue(),'application/ms-word')
            mail.send(msg)
            return render(request,'text.html',{'mail':m_id,'name':n_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':j_id,'exp':e_id,'ind':i_id,'func':f_id,'key':k_id})
        except BadHeaderError:
            return HttpResponse('Invalid header found.')
        return render(request,'text.html',{'objects_id':z_id,'mail':m_id,'name':n_id,'qua':q_id,'univ':u_id,'year':y_id,'perc':p_id,'job':j_id,'exp':e_id,'ind':i_id,'func':f_id,'key':k_id})
    
    else:
        return HttpResponse('hai')
    
def render_to_pdf(template_src, context_dict):
    template = get_template(template_src)
    context = Context(context_dict)
    html  = template.render(context)
    result = StringIO.StringIO()

    pdf = pisa.pisaDocument(StringIO.StringIO(html.encode("ISO-8859-1")), result)
    if not pdf.err:
        return HttpResponse(result.getvalue(), mimetype='application/pdf')
    return HttpResponse('We had some errors<pre>%s</pre>' % escape(html))

def myview(request):
    
    #Retrieve data or whatever you need
    return render_to_pdf(
	    #'127.0.0.1:8000/',
            'sample.html',
	    #'education.html',
	    #'professional.html',
            {
		'Full_name':'fullname',
		'Email_id':'emailid',
	            }
        )



